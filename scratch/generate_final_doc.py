from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_budget_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Official Election Infrastructure Budget Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph()
    p.add_run('Project: ').bold = True
    p.add_run('VoteGuard / JDPC Election Monitoring System\n')
    p.add_run('Status: ').bold = True
    p.add_run('Finalized Allocation for Immediate Approval')

    # Presidential Election Section
    doc.add_heading('1. 2027 National Presidential Election Budget ($8,500)', level=1)
    doc.add_paragraph('Scale: 1,000,000 Observers | 176,846 Polling Units | National Situation Room')

    table_pres = doc.add_table(rows=1, cols=2)
    table_pres.style = 'Table Grid'
    hdr_cells = table_pres.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Allocated Cost (USD)'

    pres_breakdown = [
        ("AI OCR Engine (Gemini 1.5 Flash - 200k Units)", "$250.00"),
        ("Cloud Infrastructure (Firebase DB/Storage/Functions)", "$1,200.00"),
        ("Security & DDoS Protection (Cloudflare Business Tier)", "$600.00"),
        ("High Availability & Bandwidth (10TB Surge Support)", "$1,500.00"),
        ("Authentication & Communications (Email/WhatsApp/SMS)", "$1,000.00"),
        ("Situation Room & GIS Interactive Maps", "$950.00"),
        ("Technical Support & 24/7 Monitoring (Election Week)", "$2,000.00"),
        ("Stress Testing & Load Balancing (Pre-Election)", "$1,000.00")
    ]

    for comp, cost in pres_breakdown:
        row = table_pres.add_row().cells
        row[0].text = comp
        row[1].text = cost

    last_row = table_pres.add_row().cells
    last_row[0].text = "TOTAL PRESIDENTIAL BUDGET"
    last_row[1].text = "$8,500.00"
    last_row[0].paragraphs[0].runs[0].font.bold = True
    last_row[1].paragraphs[0].runs[0].font.bold = True

    # State Election Section
    doc.add_heading('2. State-Level Election Budget ($2,300 per State)', level=1)
    doc.add_paragraph('Applies to Osun and upcoming state elections. Scale: ~5,000 Observers | Localized Results.')

    table_state = doc.add_table(rows=1, cols=2)
    table_state.style = 'Table Grid'
    hdr_cells_s = table_state.rows[0].cells
    hdr_cells_s[0].text = 'Component'
    hdr_cells_s[1].text = 'Allocated Cost (USD)'

    state_breakdown = [
        ("AI OCR Processing (State PU Results)", "$100.00"),
        ("Cloud Hosting & Database (Regional)", "$450.00"),
        ("Digital Maps & Localized GIS Dashboard", "$350.00"),
        ("Observer Authentication & Data Security", "$300.00"),
        ("Performance Monitoring & Error Tracking", "$300.00"),
        ("Technical Support (Election Day)", "$500.00"),
        ("Contingency & Surge Buffer", "$300.00")
    ]

    for comp, cost in state_breakdown:
        row = table_state.add_row().cells
        row[0].text = comp
        row[1].text = cost

    last_row_s = table_state.add_row().cells
    last_row_s[0].text = "TOTAL PER-STATE BUDGET"
    last_row_s[1].text = "$2,300.00"
    last_row_s[0].paragraphs[0].runs[0].font.bold = True
    last_row_s[1].paragraphs[0].runs[0].font.bold = True

    # Note on Cost Optimization
    doc.add_heading('3. Cost Control Strategy', level=1)
    doc.add_paragraph('• We have eliminated the $40,000 SMS wastage by implementing a hybrid authentication model.')
    doc.add_paragraph('• Gemini 1.5 Flash provides high-speed, low-cost extraction for all polling unit forms.')
    doc.add_paragraph('• The Situation Room will be optimized with Redis caching to ensure 1,000,000 users can view live results without crashing the database.')

    doc.save('Official_Election_Budget_Final.docx')
    print('Document saved as Official_Election_Budget_Final.docx')

if __name__ == "__main__":
    create_final_budget_doc()

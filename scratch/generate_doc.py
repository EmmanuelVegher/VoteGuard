from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_budget_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Election Infrastructure Budget Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph()
    p.add_run('Project: ').bold = True
    p.add_run('VoteGuard / JDPC Election App Deployment\n')
    p.add_run('Target: ').bold = True
    p.add_run('State and National Elections (2026-2027)')

    doc.add_heading('1. Smart Optical T Scanner Cost Analysis (Gemini 1.5 Flash)', level=1)
    doc.add_paragraph('By using Gemini 1.5 Flash, we eliminate the high costs of legacy OCR providers. Our system processes images as tokens, making it the most cost-effective solution for mass data entry.')

    # OCR Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Calculation'
    hdr_cells[2].text = 'Unit Cost'
    hdr_cells[3].text = 'Total (200k Units)'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Input Image Processing'
    row_cells[1].text = '200,000 x 2,000 tokens'
    row_cells[2].text = '$0.075 / 1M tokens'
    row_cells[3].text = '$30.00'

    row_cells = table.add_row().cells
    row_cells[0].text = 'JSON Output Data'
    row_cells[1].text = '200,000 x 500 tokens'
    row_cells[2].text = '$0.30 / 1M tokens'
    row_cells[3].text = '$30.00'

    doc.add_heading('2. Election Scenarios & Infrastructure Budget', level=1)

    # Scenario Table
    scenarios = [
        ("Osun State Election", "3,763", "5,000", "$450.00"),
        ("Upcoming State (Next Month)", "4,500", "6,000", "$580.00"),
        ("2027 Presidential Election", "200,000", "1,000,000", "$3,850.00")
    ]

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Scenario'
    hdr_cells2[1].text = 'Polling Units'
    hdr_cells2[2].text = 'Observers'
    hdr_cells2[3].text = 'Total Budget'

    for scenario, pus, obs, budget in scenarios:
        row = table2.add_row().cells
        row[0].text = scenario
        row[1].text = pus
        row[2].text = obs
        row[3].text = budget

    doc.add_heading('3. Detailed Breakdown (2027 Presidential)', level=1)
    doc.add_paragraph('Infrastructure cost for 1,000,000 observers reporting 200,000 polling unit results.')

    breakdown = [
        ("Smart Optical T Scanner (API)", "$60.00"),
        ("Firebase Firestore (Database)", "$850.00"),
        ("Firebase Cloud Storage (Images)", "$120.00"),
        ("Authentication (Hybrid Model)", "$1,500.00"),
        ("Hosting (Admin/Web/Situation Room)", "$350.00"),
        ("CDN & Caching (Peak Day Load)", "$470.00"),
        ("Maps & GIS Visualization", "$500.00")
    ]

    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Component'
    hdr_cells3[1].text = 'Cost Estimate'

    for comp, cost in breakdown:
        row = table3.add_row().cells
        row[0].text = comp
        row[1].text = cost

    doc.add_heading('4. Cost Optimization Strategy', level=1)
    doc.add_paragraph('• Hybrid Auth: Bulk observers use Email/Password to save $20k+ in SMS fees.')
    doc.add_paragraph('• Edge Caching: Situation Room data is cached to prevent redundant database reads.')
    doc.add_paragraph('• Serverless: No monthly server fees; pay only for actual traffic on Election Day.')

    doc.save('Election_Budget_Proposal.docx')
    print('Document saved as Election_Budget_Proposal.docx')

if __name__ == "__main__":
    create_budget_doc()
